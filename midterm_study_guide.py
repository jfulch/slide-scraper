#!/usr/bin/env python3
"""
Comprehensive Midterm Study Guide Generator

This script analyzes all slides across all lectures to create a comprehensive
midterm study guide. It extracts text, identifies key concepts, formulas, and
important topics using AI to ensure nothing is missed for exam preparation.

Usage:
    python midterm_study_guide.py
"""

import os
import json
import re
from pathlib import Path
from collections import defaultdict
import ollama
from PIL import Image
import pytesseract
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

class MidtermStudyGuideGenerator:
    def __init__(self, slides_dir="slides", output_dir="midterm_guide"):
        self.slides_dir = Path(slides_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Storage for all extracted content
        self.all_content = {}
        self.key_concepts = []
        self.formulas = []
        self.definitions = []
        self.algorithms = []
        self.examples = []
        
        print("🎓 Midterm Study Guide Generator initialized")
        print(f"📁 Slides directory: {self.slides_dir}")
        print(f"📁 Output directory: {self.output_dir}")

    def get_image_files(self, folder_path):
        """Get all image files from a folder, properly sorted"""
        image_files = []
        for ext in ['*.png', '*.jpg', '*.jpeg', '*.gif', '*.bmp']:
            image_files.extend(folder_path.glob(ext))
        
        def sort_key(path):
            filename = path.stem
            # Try to extract number from filename
            numbers = re.findall(r'\d+', filename)
            if numbers:
                return (int(numbers[0]), filename)
            return (float('inf'), filename)
        
        return sorted(image_files, key=sort_key)

    def extract_text_from_image(self, image_path):
        """Extract text from image using OCR"""
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image, config='--psm 6')
            return self.clean_text(text)
        except Exception as e:
            print(f"❌ Error extracting text from {image_path}: {e}")
            return ""

    def clean_text(self, text):
        """Clean extracted text"""
        if not text:
            return ""
        
        # Remove common OCR artifacts and university branding
        cleaning_patterns = [
            r'University of Southern California.*?(?=\n|$)',
            r'USC.*?(?=\n|$)',
            r'CSCI.*?572.*?(?=\n|$)',
            r'Information Retrieval.*?(?=\n|$)',
            r'©.*?(?=\n|$)',
            r'\b[A-Za-z]\b(?=\s|$)',  # Single letters
            r'^\s*\d+\s*$',  # Lines with only numbers
            r'^\s*[|\\/_\-=+*#@$%^&(){}[\]<>?.,;:!"\'`~]\s*$',  # Lines with only symbols
        ]
        
        cleaned_text = text
        for pattern in cleaning_patterns:
            cleaned_text = re.sub(pattern, '', cleaned_text, flags=re.IGNORECASE | re.MULTILINE)
        
        # Remove extra whitespace and empty lines
        lines = [line.strip() for line in cleaned_text.split('\n') if line.strip()]
        return '\n'.join(lines)

    def analyze_content_with_ai(self, text, lecture_name, slide_number):
        """Use AI to analyze content and extract key study materials"""
        if not text.strip():
            return None
            
        prompt = f"""
        Analyze this slide content from lecture "{lecture_name}" (slide {slide_number}) for midterm exam preparation.

        Content:
        {text}

        Please extract and categorize the following for studying:

        1. KEY CONCEPTS: Main ideas, theories, principles (mark with [CONCEPT])
        2. DEFINITIONS: Important terms and their meanings (mark with [DEFINITION])  
        3. FORMULAS/EQUATIONS: Mathematical formulas, algorithms, calculations (mark with [FORMULA])
        4. ALGORITHMS: Step-by-step processes, procedures (mark with [ALGORITHM])
        5. EXAMPLES: Concrete examples, case studies (mark with [EXAMPLE])
        6. STUDY PRIORITY: Rate importance for midterm (HIGH/MEDIUM/LOW) (mark with [PRIORITY])

        Focus on content that would likely appear on an exam. Be comprehensive but concise.
        If there are any mathematical formulas, preserve them exactly as written.
        """

        try:
            response = ollama.chat(model='llama3.1:8b', messages=[
                {'role': 'user', 'content': prompt}
            ])
            return response['message']['content']
        except Exception as e:
            print(f"❌ AI analysis error for {lecture_name} slide {slide_number}: {e}")
            return None

    def parse_ai_response(self, ai_response, lecture_name):
        """Parse AI response and categorize content"""
        if not ai_response:
            return
            
        content = {
            'lecture': lecture_name,
            'concepts': [],
            'definitions': [],
            'formulas': [],
            'algorithms': [],
            'examples': [],
            'priority': 'MEDIUM'
        }
        
        lines = ai_response.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check for section markers
            if '[CONCEPT]' in line:
                concept = line.replace('[CONCEPT]', '').strip()
                if concept:
                    content['concepts'].append(concept)
                    self.key_concepts.append({'text': concept, 'lecture': lecture_name})
                    
            elif '[DEFINITION]' in line:
                definition = line.replace('[DEFINITION]', '').strip()
                if definition:
                    content['definitions'].append(definition)
                    self.definitions.append({'text': definition, 'lecture': lecture_name})
                    
            elif '[FORMULA]' in line:
                formula = line.replace('[FORMULA]', '').strip()
                if formula:
                    content['formulas'].append(formula)
                    self.formulas.append({'text': formula, 'lecture': lecture_name})
                    
            elif '[ALGORITHM]' in line:
                algorithm = line.replace('[ALGORITHM]', '').strip()
                if algorithm:
                    content['algorithms'].append(algorithm)
                    self.algorithms.append({'text': algorithm, 'lecture': lecture_name})
                    
            elif '[EXAMPLE]' in line:
                example = line.replace('[EXAMPLE]', '').strip()
                if example:
                    content['examples'].append(example)
                    self.examples.append({'text': example, 'lecture': lecture_name})
                    
            elif '[PRIORITY]' in line:
                priority = line.replace('[PRIORITY]', '').strip().upper()
                if priority in ['HIGH', 'MEDIUM', 'LOW']:
                    content['priority'] = priority
        
        return content

    def process_all_lectures(self):
        """Process all lectures and extract content"""
        print("\n🔍 Processing all lectures for midterm study guide...")
        
        lecture_folders = [d for d in self.slides_dir.iterdir() 
                          if d.is_dir() and not d.name.startswith('.')]
        
        total_slides = 0
        processed_slides = 0
        
        for lecture_folder in sorted(lecture_folders):
            print(f"\n📚 Processing lecture: {lecture_folder.name}")
            
            image_files = self.get_image_files(lecture_folder)
            if not image_files:
                print(f"   ⚠️  No images found in {lecture_folder.name}")
                continue
                
            print(f"   📄 Found {len(image_files)} slides")
            total_slides += len(image_files)
            
            lecture_content = []
            
            for i, image_path in enumerate(image_files, 1):
                print(f"   🔎 Analyzing slide {i}/{len(image_files)}: {image_path.name}")
                
                # Extract text from image
                text = self.extract_text_from_image(image_path)
                
                if text.strip():
                    # Analyze with AI
                    ai_analysis = self.analyze_content_with_ai(text, lecture_folder.name, i)
                    parsed_content = self.parse_ai_response(ai_analysis, lecture_folder.name)
                    
                    if parsed_content:
                        parsed_content['slide_number'] = i
                        parsed_content['slide_file'] = image_path.name
                        parsed_content['raw_text'] = text
                        lecture_content.append(parsed_content)
                        processed_slides += 1
                
            self.all_content[lecture_folder.name] = lecture_content
            
        print(f"\n✅ Processing complete!")
        print(f"   📊 Total slides processed: {processed_slides}/{total_slides}")
        print(f"   🎯 Key concepts found: {len(self.key_concepts)}")
        print(f"   📝 Definitions found: {len(self.definitions)}")
        print(f"   🔢 Formulas found: {len(self.formulas)}")
        print(f"   ⚙️  Algorithms found: {len(self.algorithms)}")

    def create_comprehensive_study_guide(self):
        """Create comprehensive study guide in multiple formats"""
        print("\n📖 Creating comprehensive midterm study guide...")
        
        # Create Word document
        self.create_word_study_guide()
        
        # Create PDF study guide
        self.create_pdf_study_guide()
        
        # Create JSON summary
        self.create_json_summary()
        
        # Create topic-based summary
        self.create_topic_summary()

    def create_word_study_guide(self):
        """Create detailed Word document study guide"""
        doc = Document()
        
        # Title
        title = doc.add_heading('CSCI 572 - Information Retrieval', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Comprehensive Midterm Study Guide', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Table of contents
        doc.add_heading('Table of Contents', level=1)
        toc_content = [
            "1. High Priority Topics",
            "2. Key Concepts by Lecture", 
            "3. Important Definitions",
            "4. Formulas & Algorithms",
            "5. Examples & Case Studies",
            "6. Study Checklist"
        ]
        for item in toc_content:
            p = doc.add_paragraph(item)
            p.style = 'List Number'
        
        doc.add_page_break()
        
        # High Priority Topics
        doc.add_heading('1. High Priority Topics', level=1)
        high_priority_items = []
        for lecture_name, content_list in self.all_content.items():
            for content in content_list:
                if content.get('priority') == 'HIGH':
                    high_priority_items.extend(content.get('concepts', []))
                    high_priority_items.extend(content.get('definitions', []))
                    high_priority_items.extend(content.get('formulas', []))
        
        if high_priority_items:
            for item in high_priority_items:
                doc.add_paragraph(f"• {item}", style='List Bullet')
        else:
            doc.add_paragraph("All topics are considered important for the midterm.")
        
        doc.add_page_break()
        
        # Key Concepts by Lecture
        doc.add_heading('2. Key Concepts by Lecture', level=1)
        for lecture_name in sorted(self.all_content.keys()):
            doc.add_heading(f'{lecture_name.replace("_", " ").title()}', level=2)
            
            lecture_concepts = []
            for content in self.all_content[lecture_name]:
                lecture_concepts.extend(content.get('concepts', []))
            
            if lecture_concepts:
                for concept in lecture_concepts:
                    doc.add_paragraph(f"• {concept}", style='List Bullet')
            else:
                doc.add_paragraph("No key concepts extracted from this lecture.")
        
        doc.add_page_break()
        
        # Important Definitions
        doc.add_heading('3. Important Definitions', level=1)
        if self.definitions:
            for definition in self.definitions:
                doc.add_paragraph(f"• {definition['text']}")
                doc.add_paragraph(f"  (Source: {definition['lecture']})", style='Intense Quote')
        
        doc.add_page_break()
        
        # Formulas & Algorithms
        doc.add_heading('4. Formulas & Algorithms', level=1)
        
        if self.formulas:
            doc.add_heading('Formulas:', level=2)
            for formula in self.formulas:
                doc.add_paragraph(f"• {formula['text']}")
                doc.add_paragraph(f"  (Source: {formula['lecture']})", style='Intense Quote')
        
        if self.algorithms:
            doc.add_heading('Algorithms:', level=2)
            for algorithm in self.algorithms:
                doc.add_paragraph(f"• {algorithm['text']}")
                doc.add_paragraph(f"  (Source: {algorithm['lecture']})", style='Intense Quote')
        
        doc.add_page_break()
        
        # Examples & Case Studies
        doc.add_heading('5. Examples & Case Studies', level=1)
        if self.examples:
            for example in self.examples:
                doc.add_paragraph(f"• {example['text']}")
                doc.add_paragraph(f"  (Source: {example['lecture']})", style='Intense Quote')
        
        doc.add_page_break()
        
        # Study Checklist
        doc.add_heading('6. Study Checklist', level=1)
        checklist_items = [
            "□ Review all key concepts from each lecture",
            "□ Memorize important definitions", 
            "□ Practice all formulas and algorithms",
            "□ Work through examples and case studies",
            "□ Understand relationships between topics",
            "□ Create your own examples for each concept",
            "□ Practice applying algorithms step-by-step",
            "□ Review high-priority topics multiple times"
        ]
        
        for item in checklist_items:
            doc.add_paragraph(item, style='List Bullet')
        
        # Save document
        output_path = self.output_dir / "midterm_study_guide_comprehensive.docx"
        doc.save(output_path)
        print(f"   📄 Word study guide saved: {output_path}")

    def create_pdf_study_guide(self):
        """Create PDF version of study guide"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 24)
        
        # Title
        pdf.cell(0, 15, 'CSCI 572 - Information Retrieval', 0, 1, 'C')
        pdf.set_font('Arial', 'B', 18)
        pdf.cell(0, 10, 'Comprehensive Midterm Study Guide', 0, 1, 'C')
        pdf.ln(10)
        
        pdf.set_font('Arial', '', 12)
        
        # Summary statistics
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Study Guide Summary:', 0, 1)
        pdf.set_font('Arial', '', 12)
        
        stats = [
            f"Total Lectures Covered: {len(self.all_content)}",
            f"Key Concepts: {len(self.key_concepts)}",
            f"Important Definitions: {len(self.definitions)}",
            f"Formulas & Algorithms: {len(self.formulas) + len(self.algorithms)}",
            f"Examples & Case Studies: {len(self.examples)}"
        ]
        
        for stat in stats:
            pdf.cell(0, 8, stat, 0, 1)
        
        pdf.ln(5)
        
        # Lectures overview
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Lectures Covered:', 0, 1)
        pdf.set_font('Arial', '', 12)
        
        for lecture_name in sorted(self.all_content.keys()):
            formatted_name = lecture_name.replace('_', ' ').title()
            slide_count = len(self.all_content[lecture_name])
            pdf.cell(0, 8, f"- {formatted_name} ({slide_count} slides analyzed)", 0, 1)
        
        pdf.add_page()
        
        # Key concepts summary
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Quick Reference - Key Concepts:', 0, 1)
        pdf.set_font('Arial', '', 10)
        
        for concept in self.key_concepts[:20]:  # Limit to avoid overflow
            # Handle long text by splitting if necessary
            concept_text = concept['text'][:80] + "..." if len(concept['text']) > 80 else concept['text']
            pdf.cell(0, 6, f"- {concept_text}", 0, 1)
        
        if len(self.key_concepts) > 20:
            pdf.cell(0, 6, f"... and {len(self.key_concepts) - 20} more concepts", 0, 1)
        
        # Save PDF
        output_path = self.output_dir / "midterm_study_guide_summary.pdf"
        pdf.output(str(output_path))
        print(f"   📄 PDF study guide saved: {output_path}")

    def create_json_summary(self):
        """Create JSON summary of all extracted content"""
        summary = {
            'metadata': {
                'total_lectures': len(self.all_content),
                'total_concepts': len(self.key_concepts),
                'total_definitions': len(self.definitions),
                'total_formulas': len(self.formulas),
                'total_algorithms': len(self.algorithms),
                'total_examples': len(self.examples),
                'lectures_covered': list(self.all_content.keys())
            },
            'content_by_lecture': self.all_content,
            'consolidated': {
                'key_concepts': self.key_concepts,
                'definitions': self.definitions,
                'formulas': self.formulas,
                'algorithms': self.algorithms,
                'examples': self.examples
            }
        }
        
        output_path = self.output_dir / "midterm_study_content.json"
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, indent=2, ensure_ascii=False)
        
        print(f"   📄 JSON summary saved: {output_path}")

    def create_topic_summary(self):
        """Create topic-based summary for quick review"""
        with open(self.output_dir / "quick_topic_review.txt", 'w', encoding='utf-8') as f:
            f.write("CSCI 572 - MIDTERM STUDY GUIDE - QUICK TOPIC REVIEW\n")
            f.write("=" * 60 + "\n\n")
            
            # Group content by topic/lecture
            for lecture_name in sorted(self.all_content.keys()):
                f.write(f"{lecture_name.replace('_', ' ').upper()}\n")
                f.write("-" * 40 + "\n")
                
                lecture_content = self.all_content[lecture_name]
                if not lecture_content:
                    f.write("No content extracted.\n\n")
                    continue
                
                # Combine all concepts for this lecture
                all_concepts = []
                all_definitions = []
                all_formulas = []
                
                for content in lecture_content:
                    all_concepts.extend(content.get('concepts', []))
                    all_definitions.extend(content.get('definitions', []))
                    all_formulas.extend(content.get('formulas', []))
                
                if all_concepts:
                    f.write("KEY CONCEPTS:\n")
                    for concept in all_concepts:
                        f.write(f"• {concept}\n")
                    f.write("\n")
                
                if all_definitions:
                    f.write("DEFINITIONS:\n")
                    for definition in all_definitions:
                        f.write(f"• {definition}\n")
                    f.write("\n")
                
                if all_formulas:
                    f.write("FORMULAS/ALGORITHMS:\n")
                    for formula in all_formulas:
                        f.write(f"• {formula}\n")
                    f.write("\n")
                
                f.write("\n")
        
        print(f"   📄 Topic summary saved: {self.output_dir / 'quick_topic_review.txt'}")

    def run(self):
        """Main execution method"""
        print("🎓 Starting Comprehensive Midterm Study Guide Generation")
        print("=" * 60)
        
        try:
            # Process all lectures
            self.process_all_lectures()
            
            # Create study materials
            self.create_comprehensive_study_guide()
            
            print(f"\n🎉 Midterm study guide generation complete!")
            print(f"📁 All files saved in: {self.output_dir}")
            print("\n📋 Generated Files:")
            print("   • midterm_study_guide_comprehensive.docx (Detailed Word document)")
            print("   • midterm_study_guide_summary.pdf (Quick reference PDF)")
            print("   • midterm_study_content.json (Complete data export)")
            print("   • quick_topic_review.txt (Text-based quick review)")
            
            # Print summary statistics
            print(f"\n📊 Study Content Summary:")
            print(f"   🎯 Total Key Concepts: {len(self.key_concepts)}")
            print(f"   📝 Important Definitions: {len(self.definitions)}")
            print(f"   🔢 Formulas & Algorithms: {len(self.formulas) + len(self.algorithms)}")
            print(f"   📚 Examples & Case Studies: {len(self.examples)}")
            print(f"   🏫 Lectures Analyzed: {len(self.all_content)}")
            
        except Exception as e:
            print(f"\n❌ Error during generation: {e}")
            raise

def main():
    generator = MidtermStudyGuideGenerator()
    generator.run()

if __name__ == "__main__":
    main()