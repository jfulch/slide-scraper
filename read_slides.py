#!/usr/bin/env python3
"""
Slide Text Extractor
Extracts text from slide images using OCR and compiles it into readable documents.

This script processes all images in a lecture folder and creates:
1. A PDF document with extracted text
2. A Word document with extracted text
3. A plain text file with extracted text

Usage:
    python read_slides.py <LECTURE_NAME> [--output-format pdf|docx|txt|all]
    
Example:
    python read_slides.py querying --output-format all
    python read_slides.py youtube --output-format pdf
"""

import os
import sys
import argparse
import re
from pathlib import Path
from typing import List, Tuple, Optional
import time

try:
    import pytesseract
    from PIL import Image
except ImportError:
    print("Missing required packages. Please install:")
    print("pip install pytesseract pillow")
    sys.exit(1)

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class SlideTextExtractor:
    def __init__(self, lecture_name: str):
        self.lecture_name = lecture_name
        self.slides_dir = Path("slides") / lecture_name
        self.output_dir = Path("output")
        self.output_dir.mkdir(exist_ok=True)
        
        # Verify lecture folder exists
        if not self.slides_dir.exists():
            raise FileNotFoundError(f"Lecture folder '{self.slides_dir}' not found")
    
    def get_image_files(self) -> List[Path]:
        """Get all image files in the lecture directory, sorted by filename."""
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}
        image_files = []
        
        for file_path in self.slides_dir.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in image_extensions:
                image_files.append(file_path)
        
        # Sort by slide number if possible, otherwise alphabetically
        def sort_key(path):
            # Try to extract slide number (e.g., s1.png -> 1)
            match = re.search(r's?(\d+)', path.stem)
            if match:
                return (0, int(match.group(1)))  # Numbered slides first
            return (1, path.name)  # Non-numbered files second, alphabetically
        
        return sorted(image_files, key=sort_key)
    
    def extract_text_from_image(self, image_path: Path) -> str:
        """Extract text from a single image using OCR."""
        try:
            # Open and process the image
            with Image.open(image_path) as img:
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Use Tesseract to extract text
                # Using basic configuration for better compatibility
                text = pytesseract.image_to_string(img)
                
                # Clean up the extracted text
                text = self.clean_text(text)
                return text
                
        except Exception as e:
            print(f"Error processing {image_path}: {e}")
            return f"[Error extracting text from {image_path.name}]"
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text.strip():
            return "[No text detected]"
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
        
        # Remove university branding and generic headers
        university_patterns = [
            r'University of Southern California\s*',
            r'USC Viterbi\s*',
            r'School of Engineering\s*',
            r'USC\s*',
            r'Copyright Ellis Horowitz[,\s\d\-]*'  # Remove copyright lines
        ]
        
        for pattern in university_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
        
        # Remove lines with only special characters or very short lines
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # Skip empty lines, very short lines, and university branding
            if (len(line) > 2 and 
                not re.match(r'^[^\w\s]*$', line) and
                not re.match(r'^(University of Southern California|USC Viterbi|School of Engineering|USC)$', line, re.IGNORECASE) and
                not re.match(r'^Copyright Ellis Horowitz', line, re.IGNORECASE)):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()
    
    def extract_all_text(self) -> List[Tuple[str, str]]:
        """Extract text from all images in the lecture folder."""
        image_files = self.get_image_files()
        
        if not image_files:
            raise FileNotFoundError(f"No image files found in '{self.slides_dir}'")
        
        print(f"Processing {len(image_files)} slides from '{self.lecture_name}' lecture...")
        
        extracted_data = []
        for i, image_path in enumerate(image_files, 1):
            print(f"Processing slide {i}/{len(image_files)}: {image_path.name}")
            text = self.extract_text_from_image(image_path)
            extracted_data.append((image_path.name, text))
            
            # Small delay to be nice to system resources
            time.sleep(0.1)
        
        return extracted_data
    
    def save_as_text(self, extracted_data: List[Tuple[str, str]]) -> Path:
        """Save extracted text as a plain text file."""
        output_file = self.output_dir / f"{self.lecture_name}_slides_text.txt"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"Extracted Text from {self.lecture_name.title()} Lecture Slides\n")
            f.write("=" * 60 + "\n\n")
            
            for slide_name, text in extracted_data:
                f.write(f"Slide: {slide_name}\n")
                f.write("-" * 30 + "\n")
                f.write(text + "\n\n")
        
        return output_file
    
    def save_as_pdf(self, extracted_data: List[Tuple[str, str]]) -> Optional[Path]:
        """Save extracted text as a PDF file."""
        if not PDF_AVAILABLE:
            print("PDF output not available. Install with: pip install fpdf2")
            return None
        
        output_file = self.output_dir / f"{self.lecture_name}_slides_text.pdf"
        
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 15)
                self.cell(0, 10, f'{self.lecture_name.title()} Lecture Slides', 0, 1, 'C')
                self.ln(10)
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
        
        pdf = PDF()
        pdf.lecture_name = self.lecture_name
        pdf.add_page()
        pdf.set_font('Arial', '', 12)
        
        for slide_name, text in extracted_data:
            # Add slide header
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, f"Slide: {slide_name}", 0, 1)
            pdf.set_font('Arial', '', 10)
            
            # Add text content
            if text.strip():
                # Split text into lines and add to PDF
                for line in text.split('\n'):
                    if line.strip():
                        # Handle long lines by wrapping
                        pdf.cell(0, 6, line.encode('latin1', 'replace').decode('latin1'), 0, 1)
            else:
                pdf.cell(0, 6, "[No text detected]", 0, 1)
            
            pdf.ln(5)  # Add space between slides
        
        pdf.output(output_file)
        return output_file
    
    def save_as_docx(self, extracted_data: List[Tuple[str, str]]) -> Optional[Path]:
        """Save extracted text as a Word document."""
        if not DOCX_AVAILABLE:
            print("Word document output not available. Install with: pip install python-docx")
            return None
        
        output_file = self.output_dir / f"{self.lecture_name}_slides_text.docx"
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{self.lecture_name.title()} Lecture Slides', 0)
        
        for slide_name, text in extracted_data:
            # Add slide heading
            doc.add_heading(f'Slide: {slide_name}', level=1)
            
            # Add text content
            if text.strip():
                doc.add_paragraph(text)
            else:
                doc.add_paragraph("[No text detected]")
            
            # Add page break except for last slide
            if slide_name != extracted_data[-1][0]:
                doc.add_page_break()
        
        doc.save(output_file)
        return output_file

def check_tesseract_installation():
    """Check if Tesseract OCR is installed and accessible."""
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False

def main():
    """Parse command line arguments and run the text extractor."""
    parser = argparse.ArgumentParser(
        description="Extract text from lecture slide images using OCR",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python read_slides.py querying
  python read_slides.py youtube --output-format pdf
  python read_slides.py querying --output-format all
        """
    )
    
    parser.add_argument("lecture_name", help="Name of the lecture folder in slides/")
    parser.add_argument(
        "--output-format", 
        choices=["txt", "pdf", "docx", "all"], 
        default="all",
        help="Output format (default: all)"
    )
    
    args = parser.parse_args()
    
    # Check if Tesseract is installed
    if not check_tesseract_installation():
        print("Error: Tesseract OCR not found!")
        print("Please install Tesseract:")
        print("  macOS: brew install tesseract")
        print("  Ubuntu/Debian: sudo apt-get install tesseract-ocr")
        print("  Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki")
        sys.exit(1)
    
    try:
        # Create extractor and process slides
        extractor = SlideTextExtractor(args.lecture_name)
        extracted_data = extractor.extract_all_text()
        
        print(f"\nGenerating output files...")
        created_files = []
        
        # Generate requested output formats
        if args.output_format in ["txt", "all"]:
            txt_file = extractor.save_as_text(extracted_data)
            created_files.append(txt_file)
            print(f"✓ Text file: {txt_file}")
        
        if args.output_format in ["pdf", "all"]:
            pdf_file = extractor.save_as_pdf(extracted_data)
            if pdf_file:
                created_files.append(pdf_file)
                print(f"✓ PDF file: {pdf_file}")
        
        if args.output_format in ["docx", "all"]:
            docx_file = extractor.save_as_docx(extracted_data)
            if docx_file:
                created_files.append(docx_file)
                print(f"✓ Word document: {docx_file}")
        
        print(f"\n🎉 Successfully processed {len(extracted_data)} slides!")
        print(f"Output files saved to: {extractor.output_dir.absolute()}")
        
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
